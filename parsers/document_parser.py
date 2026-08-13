"""
Document parser: converts PDF, DOCX, and TXT files to plain text.
Also exposes a higher-level 'analyze_document' function used by the API.
"""
from __future__ import annotations
import io
from pathlib import Path
from typing import Union


# ---------------------------------------------------------------------------
# Low-level file-to-text converters
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using pdfplumber (layout-aware)."""
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _parse_pdf_bytes(data: bytes) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    from docx import Document
    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_docx_bytes(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_file(path: Union[str, Path]) -> str:
    """Parse a file on disk and return its full text content."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in {".txt", ".md"}:
        return _parse_txt(path)
    raise ValueError(f"Unsupported file extension: {ext}")


def parse_bytes(filename: str, data: bytes) -> str:
    """Parse an uploaded file (bytes + filename) and return text content."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf_bytes(data)
    if ext == ".docx":
        return _parse_docx_bytes(data)
    if ext in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")
